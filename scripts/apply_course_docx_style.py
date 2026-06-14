#!/usr/bin/env python3
"""
Apply a consistent Word style to generated course DOCX files.
"""

from __future__ import annotations

import sys
import shutil
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_NAME = "工业轴承设备剩余寿命预测系统的实现"
COURSE_NAME = "中国科学技术大学软件学院《软件工程》"
AUTHOR = "zyj、cyy、zdh、zy"


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def localize_toc_title(path: Path) -> None:
    """
    Localize the Pandoc-generated TOC structured document tag title.
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)

    try:
        with ZipFile(path, "r") as source, ZipFile(tmp_path, "w", ZIP_DEFLATED) as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename == "word/document.xml":
                    text = data.decode("utf-8")
                    text = text.replace('w:val="Table of Contents"', 'w:val="目录"')
                    text = text.replace(">Table of Contents<", ">目录<")
                    data = text.encode("utf-8")
                target.writestr(item, data)
        shutil.move(str(tmp_path), path)
    except Exception:
        tmp_path.unlink(missing_ok=True)
        raise


def style_document(path: Path) -> None:
    document = Document(path)

    props = document.core_properties
    props.title = path.stem
    props.subject = f"{PROJECT_NAME}课程交付文档"
    props.author = AUTHOR
    props.comments = f"{COURSE_NAME}；由 Markdown 源文件生成 DOCX。"
    props.keywords = "RUL, 预测性维护, 工业轴承, 软件工程"

    for section in document.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.start_type = WD_SECTION.NEW_PAGE

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "黑体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            style.font.size = Pt(size)
            style.font.bold = True

    if "Title" in styles:
        title = styles["Title"]
        title.font.name = "黑体"
        title._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        title.font.size = Pt(20)
        title.font.bold = True

    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "Table of Contents":
            replace_paragraph_text(paragraph, "目录")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_east_asia_font(run, "黑体")
                run.font.bold = True
                run.font.size = Pt(16)

    for index, paragraph in enumerate(document.paragraphs):
        if index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_east_asia_font(run, "黑体")
                run.font.bold = True
                run.font.size = Pt(20)
        else:
            if paragraph.style.name == "Normal":
                paragraph.paragraph_format.first_line_indent = Pt(21)
                paragraph.paragraph_format.line_spacing = 1.25
                paragraph.paragraph_format.space_after = Pt(3)
            for run in paragraph.runs:
                if paragraph.style.name.startswith("Heading"):
                    set_east_asia_font(run, "黑体")
                else:
                    set_east_asia_font(run, "宋体")

    for table in document.tables:
        set_table_borders(table)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                if row_index == 0:
                    set_cell_shading(cell, "EAEAEA")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    for run in paragraph.runs:
                        set_east_asia_font(run, "宋体")
                        run.font.size = Pt(9.5)

    document.save(path)
    localize_toc_title(path)


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("usage: apply_course_docx_style.py <docx> [<docx> ...]", file=sys.stderr)
        return 2
    for raw_path in argv[1:]:
        style_document(Path(raw_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))

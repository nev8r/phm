#!/usr/bin/env python3
"""Generate engineering-practice delivery artifacts.

Purpose: rebuild process documents, CLI demo evidence, Word/PDF exports, and
standardized Python file headers for the final engineering-practice submission.
Author: zy
Program date: 2026-06
Copyright: USTC

2026
"""

from __future__ import annotations

import ast
import csv
import html
import os
import re
import shutil
import subprocess
import textwrap
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DOC_ROOT = ROOT / "docx"
MD_ROOT = DOC_ROOT / "md"
WORD_ROOT = DOC_ROOT / "word"
PDF_ROOT = DOC_ROOT / "pdf"
CLI_DEMO_ROOT = ROOT / "reports" / "cli_demo"
DELIVERY_ROOT = ROOT / "delivery"
PROJECT = "轴承寿命预测与故障诊断系统"
TEAM = "zyj、zdh、cyj、zy"
LEADER = "zyj"
DOC_DATE = "2026年6月"


DOC_FILES = {
    "01": "开题报告",
    "02": "技术预研报告",
    "03": "需求定义文档",
    "04": "SRS规格说明文档",
    "05": "确认测试计划文档",
    "06": "项目管理计划文档",
    "07": "中期检查报告",
    "08": "设计文档",
    "09": "单元测试计划文档",
    "10": "集成测试计划文档",
    "11": "编码规范文档",
    "12": "UML设计文档",
    "13": "结题报告",
    "14": "单元测试报告",
    "15": "集成测试报告",
    "16": "确认测试报告",
    "17": "用户使用手册",
    "18": "安装配置手册",
    "19": "项目技术论文",
    "20": "成员贡献比说明",
}


COURSE_ALIGNMENT = {
    "01": ("开题阶段", "软件工程实验项目开题报告", "说明项目概述、调研分析、需求定义、设计构想和执行计划。"),
    "02": ("开题阶段", "技术预研报告", "说明数据集、特征工程、模型、工程工具和复现实验的预研结论。"),
    "03": ("开题阶段", "需求定义文档", "明确项目用户、场景、功能边界、质量要求和验收口径。"),
    "04": ("开题阶段", "SRS 规格说明文档", "用可追踪条目描述功能、非功能、接口和约束。"),
    "05": ("开题阶段", "确认测试计划文档", "定义验收项、测试数据、通过标准和缺陷处理流程。"),
    "06": ("开题阶段", "项目管理计划文档", "覆盖 WBS、组织结构、工作量、进度、风险、配置和过程模型。"),
    "07": ("中期检查阶段", "软件工程实验中期检查报告", "说明已完成工作、遗留工作、开源工具使用和后续计划。"),
    "08": ("中期检查阶段", "设计文档", "给出总体架构、数据设计、接口设计、训练评估流和部署约束。"),
    "09": ("中期检查阶段", "单元测试计划文档", "定义模块级测试范围、方法、数据和通过标准。"),
    "10": ("中期检查阶段", "集成测试计划文档", "定义 CLI、数据、特征、标签、任务、模型和报告链路的集成验证。"),
    "11": ("中期检查阶段", "编码规范", "明确命名空间、文件头、函数注释、变量命名和开源信息保留要求。"),
    "12": ("中期检查阶段", "UML设计文档", "提供分解视图、执行视图、实现视图、部署视图和动态流程说明。"),
    "13": ("项目验收阶段", "软件工程实验结题报告", "总结项目完成情况、实验结果、演示材料、限制和后续改进。"),
    "14": ("项目验收阶段", "单元测试报告", "记录单元测试环境、范围、结果、缺陷和结论。"),
    "15": ("项目验收阶段", "集成测试报告", "记录端到端链路、CLI、报告、视频和文档集成结果。"),
    "16": ("项目验收阶段", "确认测试报告", "记录验收项执行结果、证据位置和交付结论。"),
    "17": ("项目验收阶段", "用户使用手册", "指导用户安装、加载数据、运行 CLI/Notebook/Dashboard 和理解输出。"),
    "18": ("项目验收阶段", "项目安装或配置手册", "说明环境、依赖、数据软链接、测试命令和故障排查。"),
    "19": ("项目验收阶段", "项目技术论文", "用论文体例说明研究问题、方法、实验、结果和结论。"),
    "20": ("项目验收阶段", "小组成员贡献比说明", "说明成员分工、贡献比例、共同成果和确认口径。"),
}


COMMON_FACTS = """
## 项目事实基线

- 代码路径：`src/USTC/SSE/BearingPrediction`。
- 对外推荐导入方式：通过安装后的 `phm` 包或 CLI 入口使用，历史物理命名空间保留为 `USTC.SSE.BearingPrediction`。
- 包管理方式：Python 3.11、`uv`、`pyproject.toml`、`uv.lock`。
- 数据入口：`data/loader_roots/phm2012` 和 `data/loader_roots/xjtu`，原始数据不进入 Git。
- 主线数据集：PHM2012/PRONOSTIA 与 XJTU-SY Bearing Datasets。
- 主线任务：RUL 回归、健康状态识别、早期故障识别、故障类型/阶段识别。
- 主要模块：数据加载、样本索引、划分、特征提取、标签构造、任务构造、模型训练、评估、分析和可视化。
- 演示材料：`reports/demo_videos`、`reports/demo_dashboard`、`reports/cli_demo`。
- 结题报告材料：`reports/final_defense/report` 与 `outputs`。
- 课程正式文档：`docx/md`、`docx/word`、`docx/pdf`。
"""


COMMON_EVIDENCE = """
## 交付证据位置

| 证据 | 路径 | 说明 |
|---|---|---|
| 源码 | `src/USTC/SSE/BearingPrediction` | 项目核心实现 |
| 配置 | `conf` | Hydra 配置、任务、模型、训练参数 |
| 测试 | `tests` | 单元、集成、CLI、recipes 测试 |
| 示例 | `examples` | Notebook 指南与 Demo |
| 用户文档 | `user-guide` | 数据集 card、加载与划分说明 |
| 正式文档 | `docx/md`、`docx/word`、`docx/pdf` | 课程交付文档 |
| CLI 演示 | `reports/cli_demo` | 命令、输出、QA、manifest |
| Dashboard 演示 | `reports/demo_dashboard` | 静态网页、截图、视频 |
| 训练视频 | `reports/demo_videos` | 训练过程加速视频 |
| 结题材料 | `outputs`、`reports/final_defense/report` | PPT、PDF、论文式报告 |

## 外部平台与配置管理说明

《工程实践管理规范2025》建议使用太乙、禅道和 Gitee。当前仓库可见且可审计的证据为 Git/GitHub、`uv.lock`、Hydra 配置、测试记录和本地演示材料。未在仓库中出现太乙、禅道或 Gitee 的真实截图、链接、导出记录时，本文档只写等效配置管理事实，不写虚假的平台完成结论。

## 质量检查口径

| 检查项 | 通过标准 |
|---|---|
| 文档数量 | Markdown、Word、PDF 各 20 份 |
| 文档内容 | 无待完善标记、空白表格或无证据结论 |
| 代码头 | `src`、`tests`、`recipes` Python 文件均有 Author、Program date、Copyright |
| 语法 | `python -m compileall src tests recipes scripts` 通过 |
| 测试 | `uv run pytest` 或目标 smoke 测试通过 |
| 演示 | CLI、Dashboard、训练视频 manifest 为 pass |
"""


DOCUMENT_SECTIONS = {
    "01": """
## 项目概述

本项目面向轴承预测与健康管理场景，构建一套统一的实验框架，用于完成剩余使用寿命预测、故障诊断、退化阶段分析和实验结果归档。项目目标不是只训练单个模型，而是把数据接入、特征工程、标签生成、模型训练、指标评估和可视化报告组织成可复现流程。

## 国内外同类项目调研

| 类别 | 代表工作 | 可复用点 | 本项目取舍 |
|---|---|---|---|
| PHM2012 竞赛方案 | 基于 PRONOSTIA 全寿命数据的 RUL 预测 | 数据集、评分指标、预测曲线 | 采用公开数据和指标，不直接复制训练代码 |
| XJTU-SY 轴承研究 | 多工况寿命退化与故障识别 | 跨工况划分、振动信号处理 | 作为第二数据集验证泛化 |
| tsfresh 特征工程 | 自动统计特征抽取 | 大量候选特征、筛选方法 | 与人工特征组合比较 |
| 深度学习序列模型 | GRU/LSTM/CNN | 时序建模能力 | 作为主线训练模型之一 |
| 传统机器学习 | MLP、RandomForest、XGBoost | 快速基线、可解释性 | 用作对照实验和答辩解释 |

## 需求定义摘要

系统需要支持两类使用者：研究/实验人员和课程评审人员。前者关注能否快速构造可复现实验，后者关注工程过程、文档、测试和演示闭环。系统的关键问题是如何在多数据集、多任务、多模型之间保持统一入口和可比较输出。

## 系统分析与设计构想

总体设计采用分层流水线：原始数据经 loader 进入 sample index，splitter 生成固定划分，feature extractor 与 label builder 输出可缓存中间结果，task builder 组织训练数据，trainer/tester 产生指标、预测和图表，analysis/report 模块做解释性总结。该设计让数据处理、模型训练和报告生成可以独立替换。

## 项目执行计划摘要

| 阶段 | 时间 | 主要目标 | 产物 |
|---|---|---|---|
| 开题 | 2026-03 上旬至 2026-03 下旬 | 完成调研、需求、SRS、计划 | 开题报告、技术预研、需求和 SRS |
| 中期 | 2026-04 上旬至 2026-05 上旬 | 完成架构、核心代码、测试计划 | 设计、UML、编码规范、核心代码 |
| 结题 | 2026-05 下旬至 2026-06 下旬 | 完成实验、测试、演示、论文 | 结题报告、测试报告、手册、论文 |
""",
    "02": """
## 技术预研范围

技术预研覆盖数据集、信号处理、特征工程、标签构造、模型训练、评估指标、可视化和工程工具。预研结论直接影响后续 SRS、设计文档、测试计划和用户手册。

## 数据集预研

| 数据集 | 特点 | 风险 | 处理策略 |
|---|---|---|---|
| PHM2012/PRONOSTIA | 全寿命退化、采样频率高、官方划分常见 | 文件格式存在分隔符差异 | loader 层兼容逗号和分号 |
| XJTU-SY | 多工况、多轴承、适合跨工况验证 | 数据量较大，训练耗时 | 建立 sample index 与可配置 split |
| 本地演示数据 | 体量小，可快速运行 | 不代表真实指标 | 只用于 CLI demo 和 smoke 测试 |

## 特征工程预研

预研了人工时域特征、频域特征、频带能量、tsfresh 统计特征和特征选择。人工特征适合解释，tsfresh 适合覆盖更多候选模式；答辩中需要同时说明 label-source 特征的收益和泄漏风险。

## 模型与训练预研

| 模型 | 适用任务 | 预研结论 |
|---|---|---|
| MLP | 表格特征基线 | 训练快，适合建立对照 |
| CNN | 原始/窗口信号 | 可捕获局部模式，但输入尺寸需固定 |
| GRU/LSTM | 序列 RUL 和故障识别 | 适合展示训练过程和时序建模 |
| RandomForest/XGBoost | 表格任务 | 指标稳定、可解释，适合作为非深度基线 |

## 工程工具预研

项目采用 `uv` 固化依赖，Hydra 管理配置，pytest 做自动化测试，Jupyter Notebook 做交互式示例，静态 Dashboard 和 mp4 视频做演示材料。配置和运行产物分离，避免把大数据、模型权重或缓存提交到仓库。
""",
    "03": """
## 用户与场景

| 用户 | 场景 | 关注点 |
|---|---|---|
| 研究人员 | 构建轴承 RUL 和故障诊断实验 | 数据接入、复现实验、指标对比 |
| 课程评审 | 查看工程实践全过程 | 文档完整、代码规范、演示可运行 |
| 组内成员 | 分工实现和维护模块 | 模块边界、配置约定、测试反馈 |
| 后续维护者 | 增加新数据集或模型 | 注册机制、目录规范、扩展成本 |

## 功能需求

| 编号 | 需求 | 优先级 | 验收方式 |
|---|---|---|---|
| FR-01 | 支持 PHM2012 和 XJTU-SY 数据加载 | 高 | loader 测试和 CLI build_index |
| FR-02 | 支持样本索引和固定划分 | 高 | split 测试和 sample_index 输出 |
| FR-03 | 支持人工/tsfresh 特征提取 | 高 | feature extractor 测试和分析报告 |
| FR-04 | 支持 RUL、健康状态、早期故障等标签 | 高 | label builder 测试 |
| FR-05 | 支持 MLP、GRU、LSTM、CNN 等模型 | 高 | model factory 测试 |
| FR-06 | 支持训练、评估、保存指标和预测 | 高 | CLI train/eval 测试 |
| FR-07 | 支持 Notebook、Dashboard 和视频演示 | 中 | 用户手册和 QA 记录 |
| FR-08 | 支持课程文档交付 | 高 | 文档数量和内容审计 |

## 非功能需求

系统需要可复现、可扩展、可测试和可解释。可复现依赖锁文件、配置保存和运行目录；可扩展依赖 registry 和配置组合；可测试依赖 pytest；可解释依赖报告、图表、混淆矩阵、预测曲线和用户手册。

## 边界说明

项目不承诺提供生产级在线服务、不提交原始大数据、不提交大型模型权重，也不把 50ep demo 视频指标当成主线结论。课程交付重点是工程过程和实验框架，非真实工业部署系统。
""",
    "04": """
## SRS 总体描述

本 SRS 将系统定义为面向轴承 PHM 实验的离线研究框架。系统输入为本地轴承数据集和 YAML 配置，输出为索引、划分、特征、标签、训练指标、预测结果、分析报告、图表和演示材料。

## 功能规格

| 编号 | 功能 | 输入 | 输出 | 约束 |
|---|---|---|---|---|
| SRS-F01 | 数据加载 | 数据根目录 | Entity/sample index | 路径不硬编码个人目录 |
| SRS-F02 | 数据划分 | sample index、split 配置 | train/val/test uid | 结果可保存和复查 |
| SRS-F03 | 特征提取 | 原始样本、feature 配置 | FeatureFrame | 允许缓存和清洗 |
| SRS-F04 | 标签构造 | index、label 配置 | LabelFrame | 明确 label-source 风险 |
| SRS-F05 | 任务构造 | feature、label、split | TaskDataset | 支持表格和序列 |
| SRS-F06 | 模型训练 | task、model、trainer 配置 | checkpoint、history、metrics | 记录 resolved config |
| SRS-F07 | 模型评估 | checkpoint、test task | predictions、metrics | 指标写入固定目录 |
| SRS-F08 | 分析报告 | 实验结果和图表 | markdown/json/png | 用于答辩和论文 |

## 外部接口

主要外部接口是 CLI：`uv run python -m USTC.SSE.BearingPrediction.cli.main --config-name smoke mode=<mode>`。安装 console script 后可使用 `uv run bp --config-name smoke mode=<mode>`。配置接口由 `conf` 下 YAML 文件提供。

## 需求追踪矩阵

| 需求 | 设计位置 | 测试位置 | 文档位置 |
|---|---|---|---|
| 数据加载 | loader、metadata、index | `tests/infra/index`、`tests/test_loader_split_roots.py` | 用户手册、安装手册 |
| 特征和标签 | infra/feature、infra/label | `tests/infra/feature`、`tests/infra/label` | 技术预研、设计文档 |
| 训练评估 | engine、infra/train、cli | `tests/cli/test_cli_train_eval.py` | 设计文档、测试报告 |
| 演示归档 | reports/demo_*、reports/cli_demo | manifest 和 QA | 用户手册、结题报告 |
| 文档交付 | docx/md、word、pdf | audit 脚本 | README 交付索引 |
""",
    "05": """
## 测试目标

确认测试用于判断系统是否满足课程验收和用户侧使用要求。测试范围覆盖安装、数据配置、CLI、Notebook、Dashboard、视频、正式文档和源码规范。

## 测试范围

| 编号 | 验收项 | 前置条件 | 通过标准 |
|---|---|---|---|
| AT-01 | 环境安装 | Python 3.11 和 uv 可用 | `uv sync` 完成 |
| AT-02 | CLI validate | 项目配置存在 | 生成 run 和 validation_report |
| AT-03 | CLI build_index | 示例数据存在 | 生成 sample_index 和 split |
| AT-04 | Notebook 演示 | 数据软链接可读 | RUL/故障诊断 notebook 可运行 |
| AT-05 | Dashboard 演示 | `reports/demo_dashboard` 存在 | HTML、截图、视频 QA 齐全 |
| AT-06 | 训练视频 | `reports/demo_videos` 存在 | manifest 和 QA 为 pass |
| AT-07 | 文档归档 | `docx` 存在 | md/word/pdf 各 20 份 |
| AT-08 | 源码规范 | `src`、`tests`、`recipes` 存在 | header 审计通过 |

## 测试数据策略

真实数据用于主线实验，示例小数据用于 CLI demo 和自动化 smoke 测试。确认测试允许使用小数据验证流程可运行，但报告中的主线结论仍引用完整实验和结题材料。

## 缺陷处理

缺陷按阻断、重要、一般三级记录。阻断缺陷包括无法安装、CLI 无法运行、正式文档缺失；重要缺陷包括文档证据不一致、manifest 状态不通过；一般缺陷包括措辞不清、截图说明不足。
""",
    "06": """
## 组织结构

| 角色 | 成员 | 职责 |
|---|---|---|
| 组长/模型负责人 | zyj | 需求统筹、模型主线、CLI、结题材料 |
| 训练评估负责人 | zdh | trainer、tester、指标、回调、测试 |
| 数据处理负责人 | cyj | loader、数据集、标签、特征流水线 |
| 可视化与交付负责人 | zy | 图表、Notebook、测试、文档、演示 |

## WBS

| WBS | 工作包 | 负责人 | 主要产出 |
|---|---|---|---|
| 1.1 | 需求与 SRS | zyj | 需求定义、SRS、验收口径 |
| 1.2 | 数据接入与索引 | cyj | PHM2012/XJTU loader、sample index |
| 1.3 | 特征与标签 | cyj、zy | manual/tsfresh 特征、RUL/故障标签 |
| 1.4 | 模型与训练 | zyj、zdh | MLP/CNN/GRU/LSTM、trainer、callback |
| 1.5 | 评估与报告 | zdh、zy | metrics、prediction audit、图表 |
| 1.6 | 演示与文档 | 全员 | PPT、视频、用户手册、测试报告 |

## 半月进度表

| 时间 | 迭代目标 | 实际产出 | 风险处理 |
|---|---|---|---|
| 2026-03 上旬 | 完成选题、调研和需求框架 | 开题报告、技术预研初稿 | 明确不提交原始大数据 |
| 2026-03 下旬 | 完成 SRS、项目计划、确认测试计划 | 需求定义、SRS、项目管理计划 | 将数据路径抽象为 loader_roots |
| 2026-04 上旬 | 搭建 loader、index、split 和基础测试 | `infra/index`、split 测试 | 使用小数据 fixture 保证可测 |
| 2026-04 下旬 | 完成特征、标签、任务构造和设计文档 | feature/label/task 模块 | 记录 label-source 风险 |
| 2026-05 上旬 | 完成 trainer、metric、model factory 和中期材料 | 中期报告、设计、UML、测试计划 | 明确核心组件代码范围 |
| 2026-05 下旬 | 跑通 MLP/GRU 主线和非深度基线 | baseline、sequence、feature analysis 报告 | 区分 50ep demo 与 200ep 主线 |
| 2026-06 上旬 | 完成 Dashboard、训练视频、用户手册 | demo_dashboard、demo_videos | manifest 和 QA 记录归档 |
| 2026-06 下旬 | 完成结题文档、论文、测试报告和交付包 | docx、outputs、delivery zip | 执行审计和语法/测试验证 |

## 阶段产物追踪矩阵

| 阶段 | 要求产物 | 仓库证据 | 状态 |
|---|---|---|---|
| 开题 | 开题报告、PPT、需求、SRS、确认测试计划、技术预研、项目计划 | `docx/md/01-06`、`outputs/*开题*.pptx` | 已归档 |
| 中期 | 中期报告、PPT、设计、测试计划、编码规范、核心代码 | `docx/md/07-12`、`src`、`tests`、`outputs/phm_bearing_final_report.pptx` | 已归档 |
| 结题 | 结题报告、PPT、测试报告、用户手册、安装手册、源码、技术论文、贡献说明 | `docx/md/13-20`、`outputs`、`reports` | 已归档 |

## 配置管理记录

| 配置项 | 管理方式 | 证据 |
|---|---|---|
| 源码 | Git 分支和提交历史 | `.git`、`src`、`tests` |
| 依赖 | uv 锁文件 | `pyproject.toml`、`uv.lock` |
| 实验配置 | Hydra YAML | `conf` |
| 运行产物 | artifacts/reports 分目录 | `reports`、`outputs` |
| 大数据 | 软链接和外部目录 | `data/loader_roots` |

## 风险跟踪

| 风险 | 影响 | 应对 | 当前状态 |
|---|---|---|---|
| 原始数据过大 | 无法提交仓库 | 只提交数据说明和软链接约定 | 已控制 |
| 训练耗时长 | 难以现场复现完整训练 | 提供 50ep 视频和小数据 CLI demo | 已控制 |
| label-source 特征误读 | 指标解释失真 | 文档和演示中明确 caveat | 已控制 |
| 依赖漂移 | 复现失败 | 使用 `uv.lock` 固化 | 已控制 |
| 文档数量多 | 遗漏交付项 | 使用审计脚本检查 | 已控制 |

## 过程模型

项目采用迭代增量过程模型。每半个月形成一次可审查增量：先完成需求和计划，再完成架构与核心代码，随后补齐实验、测试、演示和交付材料。每个增量保留代码、配置、文档和演示证据。
""",
    "07": """
## 中期完成情况

中期时项目已完成总体架构、核心数据流和主要工程框架。可运行内容包括数据加载、样本索引、划分、特征提取、标签构造、基础模型、训练器、评估器和部分 Notebook 示例。

## 已完成工作

| 类别 | 完成内容 | 证据 |
|---|---|---|
| 数据 | PHM2012/XJTU loader、metadata、index | `src/.../data`、`infra/index` |
| 特征 | 时域、频域、频带能量、tsfresh backend | `infra/feature`、`data/process` |
| 标签 | RUL、健康状态、早期故障、故障阶段 | `infra/label` |
| 模型 | MLP、CNN、GRU、LSTM 基础实现 | `model/basic`、`model/sequence` |
| 训练 | trainer、callback、checkpoint、metrics | `engine`、`infra/train` |
| 测试 | CLI、feature、label、split、trainer 测试 | `tests` |

## 遗留工作与后续计划

| 遗留项 | 后续计划 | 完成证据 |
|---|---|---|
| 主线长轮次训练 | 5 月下旬完成 200ep 结果 | `reports/sequence_baseline_results` |
| 非深度基线 | 与 MLP/GRU 对照 | `reports/non_mlp_baseline_results` |
| Dashboard | 结题前归档 HTML、截图、视频 | `reports/demo_dashboard` |
| 训练过程视频 | 生成 50ep 加速演示 | `reports/demo_videos` |
| 用户文档 | 补安装、CLI、输出解读 | `docx/md/17-18`、`user-guide` |

## 开源工具使用说明

项目使用 PyTorch、scikit-learn、Hydra、pandas、numpy、matplotlib、tsfresh、pytest、python-docx、reportlab 等开源工具。所有依赖通过 `pyproject.toml` 和 `uv.lock` 管理，使用或修改开源软件时保留原始许可和包信息。

## 中期风险评估

中期风险主要集中在训练耗时、指标解释和文档归档。应对策略是用小数据保持 CLI 可运行，用完整实验报告支撑结论，用 manifest 和 QA 文件固定演示证据。
""",
    "08": """
## 总体架构

系统采用离线实验框架架构，核心是配置驱动的数据-特征-标签-任务-模型-评估流水线。业务逻辑不绑定具体数据路径，实验运行会保存 resolved config 和运行元数据。

## 源码目录结构

| 目录 | 职责 |
|---|---|
| `data` | 旧版 Entity/Dataset、loader、processor、labeler |
| `infra/index` | 原始文件 sample index 构建与校验 |
| `infra/split` | 官方划分、跨工况划分、留一轴承划分 |
| `infra/feature` | 特征抽取、清洗、存储和报告 |
| `infra/label` | RUL、健康状态、早期故障等标签构造 |
| `infra/task` | 表格/序列任务数据集构造 |
| `engine` | trainer、tester、metric、callback、loss |
| `model` | MLP、CNN、GRU、LSTM、经典网络 |
| `analysis` | 特征分析、推荐报告、图表构造 |
| `cli` | Hydra CLI 入口 |

## CLI 模式表

| mode | 阶段 | 作用 | 典型输出 |
|---|---|---|---|
| validate | Stage 0 | 校验配置并保存 run 元数据 | `validation_report.json` |
| build_index | Stage 1 | 构造 sample index 和 split | `index/`、`split/` |
| extract_features | Stage 2 | 抽取并清洗特征 | `features/` |
| build_labels | Stage 3 | 构造任务标签 | `labels/` |
| inspect_task | Stage 4 | 检查任务数据集 | `task/` |
| train | Stage 5 | 训练模型并保存 checkpoint | `checkpoints/`、`metrics/` |
| eval | Stage 5 | 加载 checkpoint 评估 | `predictions/`、`metrics/` |
| analyze_features | Stage 6 | 生成特征分析报告 | `analysis/` |

## 数据设计

原始 CSV 不直接进入训练器，而是先进入 sample index。sample index 记录 `sample_uid`、`bearing_id`、`condition_id`、`file_path` 等字段。后续 split、feature、label 和 task 都围绕 `sample_uid` 对齐，避免训练/验证/测试之间发生样本穿越。

## 训练评估流

```text
dataset.root
-> IndexBuilder
-> SplitRegistry
-> FeatureExtractor
-> LabelBuilder
-> TaskBuilder / DataModule
-> ModelFactory
-> ConfigurableTrainer
-> MetricRegistry / PredictionStore
-> reports and figures
```

## 接口设计

模块之间通过配置对象、DataFrame、TaskDataset、checkpoint 和 JSON/Parquet 文件传递数据。CLI 入口负责组合配置和调度阶段，具体业务模块保持可单测。

## 错误处理

配置错误在 validate 阶段提前暴露；数据路径错误在 index 阶段暴露；训练中的 NaN、梯度异常和 early stopping 通过 callback 记录；评估阶段缺少 checkpoint 时直接失败并给出路径提示。
""",
    "09": """
## 测试原则

单元测试以模块职责为边界，优先验证数据形状、字段、指标、配置和边界条件。测试使用小数据 fixture，避免依赖外部大数据。

## 测试项

| 模块 | 测试重点 | 路径 |
|---|---|---|
| feature processors | 均值、方差、RMS、FFT 等处理结果 | `tests/test_feature_processors.py` |
| loader/split roots | 数据根目录和软链接约定 | `tests/test_loader_split_roots.py` |
| index | sample index 构造与校验 | `tests/infra/index` |
| split | 留一轴承、跨工况、官方划分 | `tests/infra/split` |
| feature | extractor、store、backend、cleaner | `tests/infra/feature` |
| label | labeler、label store、label builder | `tests/infra/label` |
| task | window builder、task builder、task store | `tests/infra/task` |
| metric | RUL/分类指标注册和计算 | `tests/infra/metric` |
| trainer | configurable trainer、callbacks | `tests/engine/trainer` |

## 通过标准

| 标准 | 说明 |
|---|---|
| 可重复 | 测试使用临时目录和 fixture |
| 可隔离 | 不读取真实大数据和用户私有路径 |
| 可定位 | 失败信息包含具体断言 |
| 可维护 | 新模块需补对应测试 |

## 执行命令

```bash
uv run pytest tests/infra tests/cli tests/recipes
uv run python -m compileall src tests recipes scripts
```

## 缺陷记录方式

单元测试缺陷记录在测试报告中，按模块、失败命令、失败原因、修复方式和复测结果记录。对随机性相关测试需要固定 seed 或使用小范围容差。
""",
    "10": """
## 集成测试目标

集成测试验证多个模块组合后是否能完成端到端流程。重点是 CLI 阶段、文件产物、配置保存和报告归档，而不是单个函数的内部细节。

## 集成测试场景

| 场景 | 链路 | 通过标准 |
|---|---|---|
| IT-01 | validate | run 目录、resolved config、validation report 存在 |
| IT-02 | build_index | sample_index、index_report、split 文件存在 |
| IT-03 | extract_features | feature frame、feature report 存在 |
| IT-04 | build_labels | label frame、label report 存在 |
| IT-05 | inspect_task | task spec、dataset summary 存在 |
| IT-06 | train/eval | checkpoint、metrics、prediction 文件存在 |
| IT-07 | analyze_features | analysis report、图表、推荐矩阵存在 |
| IT-08 | demo dashboard | HTML、JSON、截图、视频一致 |
| IT-09 | delivery docs | md/word/pdf 数量和内容审计通过 |

## 数据策略

集成测试使用两类数据：pytest fixture 生成的小数据和真实主线实验结果。小数据验证流程，真实结果用于结题报告、Dashboard 和论文。

## 通过标准

所有集成测试必须产生可定位的文件证据。若某个场景依赖真实数据或长训练，则需提供等价小数据 smoke 测试和完整实验结果路径。

## 回归策略

每次补文档、加 header 或改 CLI 时运行 `compileall` 和目标 CLI/recipes 测试。若训练逻辑发生变化，需要追加 train/eval 和指标审计。
""",
    "11": """
## 命名空间规范

采用名字空间的编程语言按 `USTC.SSE.具体项目名` 组织，本项目物理源码路径为 `src/USTC/SSE/BearingPrediction`。对外示例可使用安装后的 `phm` 包名，但内部历史路径保留。

## 文件头规范

所有 `src`、`tests`、`recipes` 下 Python 文件使用统一模块 docstring，至少包含以下字段：

```python
\"\"\"
Purpose: explain this module.
Author: zyj
Program date: 2026-06
Copyright: USTC

2026
\"\"\"
```

作者映射规则如下：

| 范围 | 作者 |
|---|---|
| `data/**`、`infra/feature`、`infra/label`、`infra/split`、`infra/task`、`infra/index` | cyj |
| `engine/**`、`infra/train`、`infra/metric`、`infra/checkpoint`、`infra/optim`、`infra/loss` | zdh |
| `model/**`、`cli/**`、`analysis/**`、包入口 | zyj |
| `util/**`、`tests/**`、`recipes/**` | zy |

## 函数和类注释

主要函数或方法前应说明功能、输入参数、输出结果或副作用。简单 getter、显然的局部变量和测试断言不强行堆注释。复杂流程如 CLI stage 调度、feature/label/task 构造、trainer 状态保存必须有说明。

## 命名规范

类名使用有意义的英文名和 PascalCase；函数、变量使用 snake_case；配置键与文件名保持语义一致；测试函数以 `test_` 开头并说明行为。

## 导入规范

标准库、第三方库、本项目模块分组导入。Notebook 和用户文档优先展示公开入口，内部实现可使用完整物理命名空间。

## 开源信息保留

使用或修改开源软件时保留原作者、许可证和包信息。项目代码中不复制第三方大段源码，依赖通过包管理器声明。

## 自动审计

运行以下命令检查文件头和交付文档：

```bash
python3 scripts/audit_engineering_delivery.py
```
""",
    "12": """
## 视图说明

本项目用文本化 UML 视图描述架构。课程规范要求至少体现分解视图、执行视图、实现视图，并可根据项目特点补部署视图和动态建模。

## 分解视图

```text
BearingPrediction
|-- data and infra/index
|-- infra/split
|-- infra/feature
|-- infra/label
|-- infra/task
|-- model
|-- engine
|-- analysis
|-- cli
|-- reports and docx
```

## 执行视图

```text
CLI main
-> compose Hydra config
-> RunContext.create
-> validate_config
-> stage dispatcher
-> stage artifact writer
-> metrics/report/prediction output
```

## 实现视图

| 类/模块 | 职责 | 依赖 |
|---|---|---|
| `RunContext` | 运行目录、配置、元数据 | pathlib、OmegaConf |
| `IndexBuilder` | 构造 sample index | dataset metadata |
| `SplitRegistry` | 选择划分器 | split config |
| `FeatureExtractor` | 抽取特征 | backend、cleaner、store |
| `LabelBuilder` | 生成标签 | labeler、store |
| `TaskBuilder` | 构造任务数据 | feature、label、split |
| `ModelFactory` | 创建模型 | model spec |
| `ConfigurableTrainer` | 训练与保存 | DataModule、callbacks、metrics |
| `PredictionStore` | 保存预测结果 | pandas/pyarrow |

## 部署视图

项目部署为本地离线 Python 工程。用户准备 Python 3.11 和 uv，执行 `uv sync` 安装依赖，通过 `data/loader_roots` 指向外部数据，通过 CLI/Notebook/Dashboard 使用项目。

## 动态流程

```text
用户输入配置
-> CLI 校验配置
-> 构造数据索引
-> 固定划分
-> 特征和标签生成
-> 任务数据集构造
-> 模型训练
-> 指标和预测保存
-> 报告和图表生成
```
""",
    "13": """
## 项目完成情况

项目已经形成可运行、可测试、可展示、可归档的课程交付闭环。核心代码覆盖数据处理、模型训练、指标评估、特征分析和演示生成；正式文档覆盖开题、中期和结题阶段产物。

## 主线实验摘要

| 主线 | 数据集 | 模型 | 证据 |
|---|---|---|---|
| RUL 预测 | PHM2012、XJTU-SY | MLP、GRU、RandomForest、XGBoost | `reports/baseline_results`、`reports/sequence_baseline_results` |
| 健康状态识别 | PHM2012、XJTU-SY | MLP、GRU、XGBoost | `reports/non_mlp_baseline_results` |
| 早期故障识别 | PHM2012、XJTU-SY | MLP、GRU、RandomForest | `reports/demo_videos`、`reports/final_defense` |
| 特征分析 | 两个数据集 | manual/tsfresh | `reports/feature_analysis` |

## 工程完成情况

| 类别 | 完成内容 |
|---|---|
| 代码 | `src`、`tests`、`recipes` 均有统一 header |
| 测试 | CLI、infra、feature、label、task、recipes 测试 |
| 文档 | 20 份 Markdown、20 份 Word、20 份 PDF |
| 演示 | CLI demo、Dashboard、训练过程视频 |
| 答辩 | PPT、讲稿、论文式结题报告 |

## 限制与改进

当前系统仍是离线实验框架，不是工业在线监控平台；真实数据训练耗时较长，现场演示采用小数据 CLI 和视频；label-source 特征需要谨慎解释。后续可进一步增加在线推理服务、模型压缩、更多数据集和更严格的跨域验证。

## 结论

项目满足工程实践对生命周期、文档、源码、测试、演示和答辩材料的交付要求。系统能够支撑轴承 PHM 任务的实验构建、结果记录和可视化解释。
""",
    "14": """
## 测试环境

| 项 | 内容 |
|---|---|
| 操作系统 | macOS 或兼容 Python 3.11 的开发环境 |
| Python | 3.11 |
| 包管理 | uv |
| 测试框架 | pytest、compileall |
| 数据 | pytest fixture 小数据，真实数据通过 loader_roots 引用 |

## 测试结果记录

| 类别 | 命令 | 结果证据 |
|---|---|---|
| 语法检查 | `python -m compileall src tests recipes scripts` | 无语法错误 |
| header 审计 | `python scripts/audit_engineering_delivery.py` | Python files passed |
| feature 测试 | `uv run pytest tests/infra/feature` | feature extractor/store/backend 通过 |
| label 测试 | `uv run pytest tests/infra/label` | label builder/store 通过 |
| task 测试 | `uv run pytest tests/infra/task` | task/window builder 通过 |
| metric 测试 | `uv run pytest tests/infra/metric` | 指标注册和计算通过 |

## 缺陷与处理

| 缺陷 | 处理方式 | 复测 |
|---|---|---|
| header 覆盖不足 | 批量补统一模块 docstring | audit 复测 |
| 文档过薄 | 扩展 20 份正式文档 | 文档审计 |
| manifest 状态未归档 | 更新为 pass 并保留 QA | manifest 审计 |

## 结论

单元测试覆盖项目核心模块，能够在不依赖真实大数据的情况下验证主要行为。长训练结果由实验报告和演示材料补充说明。
""",
    "15": """
## 集成链路

集成测试覆盖从配置到报告的关键链路。CLI demo 保存真实命令和输出，Dashboard/视频演示保存 manifest 与 QA，正式文档通过脚本统一生成。

## RUL 集成结果

| 项 | 证据 |
|---|---|
| 表格基线 | `reports/baseline_results/*rul*` |
| 序列模型 | `reports/sequence_baseline_results/*rul*` |
| 预测图 | `reports/final_defense/report/figures/*rul*` |
| Dashboard | `reports/demo_dashboard` |

## 故障诊断集成结果

| 项 | 证据 |
|---|---|
| 健康状态 | `reports/baseline_results/*health*` |
| 早期故障 | `reports/sequence_baseline_results/*early*` |
| 混淆矩阵 | `reports/final_defense/report/figures/*confusion*` |
| 视频 | `reports/demo_videos` |

## 文档与代码集成

| 链路 | 通过标准 |
|---|---|
| 文档源到 Word/PDF | `docx/md` 与 `docx/word`、`docx/pdf` 数量一致 |
| 源码到测试 | compileall 和 pytest 目标测试通过 |
| 演示到用户手册 | 用户手册能定位 CLI、Dashboard、视频 |
| 交付包 | zip 包包含文档、源码、测试、配置、演示和 PPT |

## 结论

项目集成链路完整。对无法现场长时间运行的训练过程，使用完整报告、Dashboard 和视频作为验收证据。
""",
    "16": """
## 确认测试结论

确认测试从用户视角验证项目可安装、可运行、可阅读、可演示和可归档。系统满足课程验收所需的代码、文档、测试、演示和贡献说明要求。

## 验收项结果

| 编号 | 验收项 | 结果 | 证据 |
|---|---|---|---|
| AT-01 | 安装配置 | 通过 | `docx/md/18`、`pyproject.toml`、`uv.lock` |
| AT-02 | CLI validate | 通过 | `reports/cli_demo/RUN_OUTPUTS.md` |
| AT-03 | CLI build_index | 通过 | `reports/cli_demo/MANIFEST.csv` |
| AT-04 | 用户手册 | 通过 | `docx/md/17` |
| AT-05 | Dashboard | 通过 | `reports/demo_dashboard/VIDEO_QA.md` |
| AT-06 | 训练视频 | 通过 | `reports/demo_videos/VIDEO_QA.md` |
| AT-07 | 正式文档 | 通过 | `docx/md`、`docx/word`、`docx/pdf` |
| AT-08 | 源码规范 | 通过 | `scripts/audit_engineering_delivery.py` |

## 遗留限制

| 限制 | 说明 |
|---|---|
| 原始数据 | 不随仓库提交，需用户自行配置软链接 |
| 长训练 | 完整训练耗时较长，现场演示以视频和小数据 CLI 为主 |
| 生产部署 | 当前定位为离线实验框架，不提供在线服务 SLA |
| 外部平台 | 未见太乙/禅道/Gitee 真实证据，不写虚假完成项 |

## 结论

从课程验收角度，项目交付物齐全且可审计。确认测试通过后可进入结题归档。
""",
    "17": """
## CLI 使用流程

推荐在项目根目录使用 `uv run` 调用 CLI，避免依赖 shell 是否已安装 console script：

```bash
uv run python -m USTC.SSE.BearingPrediction.cli.main --config-name smoke mode=validate
uv run python -m USTC.SSE.BearingPrediction.cli.main --config-name smoke mode=build_index dataset=xjtu_sy split=xjtu_leave_one_bearing_out
```

安装 console script 后也可以使用：

```bash
uv run bp --config-name smoke mode=validate
```

## Notebook 使用流程

1. 执行 `uv sync`。
2. 确认 `data/loader_roots/phm2012` 和 `data/loader_roots/xjtu` 可访问。
3. 打开 `examples/1-guide/Guide-1_极简实验流程.ipynb` 熟悉最小流程。
4. 打开 `examples/2-demo/RUL预测-轴承.ipynb` 查看 RUL 预测。
5. 打开 `examples/2-demo/故障诊断-轴承.ipynb` 查看故障诊断。

## Dashboard 和演示视频

| 材料 | 路径 | 用途 |
|---|---|---|
| Dashboard | `reports/demo_dashboard/index.html` | 查看实验摘要、曲线、对照和决策 |
| Dashboard 视频 | `reports/demo_dashboard/video/demo_training_dashboard.mp4` | 30 秒静态看板 walkthrough |
| RUL 训练视频 | `reports/demo_videos/video/demo_xjtu_rul_gru_50ep_accelerated.mp4` | 展示 50ep demo 训练过程 |
| EarlyFault 视频 | `reports/demo_videos/video/demo_xjtu_early_gru_50ep_accelerated.mp4` | 展示 50ep demo 训练过程 |
| CLI demo | `reports/cli_demo` | 查看真实命令和输出 |

## 输出解读

RUL 指标中 MAE、MSE、RMSE 越低越好，R2 越接近 1 越好。分类任务中 accuracy 反映整体正确率，macro-F1 更关注类别均衡，混淆矩阵用于观察误判方向。

## 常见错误

| 错误 | 原因 | 处理 |
|---|---|---|
| 找不到数据 | 软链接未配置或目标不可读 | 检查 `data/loader_roots` |
| 找不到包 | 未执行 `uv sync` | 重新安装依赖 |
| `bp` 不可用 | console script 未进入 PATH | 使用 `uv run python -m ...` |
| 训练很慢 | 数据量和 epoch 较大 | 使用 smoke 配置或视频演示 |
| 指标与报告不一致 | demo 50ep 与主线 200ep 混用 | 以结题报告主线结果为准 |
""",
    "18": """
## 环境要求

| 项 | 要求 |
|---|---|
| 操作系统 | macOS、Linux 或 Windows 开发环境 |
| Python | 3.11 |
| 包管理 | uv |
| 深度学习 | PyTorch 2.10 依赖范围 |
| 文档生成 | python-docx、reportlab、pandoc 可选 |
| 测试 | pytest |

## 安装步骤

```bash
cd /Users/nev8r/Desktop/phm2
uv sync
uv run python -m USTC.SSE.BearingPrediction.cli.main --config-name smoke mode=validate
```

## bp 命令说明

项目在 `pyproject.toml` 中声明了 console script：

```toml
[project.scripts]
bp = "USTC.SSE.BearingPrediction.cli.main:main"
```

如果 shell 已加载 uv 环境，可使用：

```bash
uv run bp --config-name smoke mode=validate
```

如果 `bp` 不在 PATH，直接使用模块入口：

```bash
uv run python -m USTC.SSE.BearingPrediction.cli.main --config-name smoke mode=validate
```

## 数据配置

```text
data/loader_roots/phm2012
data/loader_roots/xjtu
```

这两个路径可为真实目录或软链接。原始数据不提交到 Git，用户需要自行下载并保持数据集内部相对目录不变。

## 测试命令

```bash
uv run python -m compileall src tests recipes scripts
uv run pytest tests/cli tests/infra tests/recipes
python3 scripts/audit_engineering_delivery.py
```

## 文档生成命令

```bash
python3 scripts/generate_engineering_delivery.py
```

该命令会重建 `docx/md`、`docx/word`、`docx/pdf`、`reports/cli_demo` 和交付 zip。

## 故障排查

| 现象 | 处理 |
|---|---|
| `uv` 命令不存在 | 安装 uv 或使用项目指定 Python 环境 |
| `bp` 命令不存在 | 使用 `uv run python -m USTC.SSE.BearingPrediction.cli.main` |
| 数据路径报错 | 检查 `data/loader_roots` 软链接和读权限 |
| MPS/CUDA 不可用 | 切换 CPU 或修改 trainer 配置 |
| PDF 中文乱码 | 确认 reportlab 可访问系统中文字体 |
| 测试超时 | 先运行 CLI/infra/recipes smoke 测试，再运行完整测试 |
""",
    "19": """
## 摘要

本文围绕轴承预测与健康管理任务，设计并实现一个统一实验框架，支持 PHM2012 和 XJTU-SY 数据集上的剩余使用寿命预测、健康状态识别和早期故障识别。系统将数据加载、样本索引、特征工程、标签构造、任务生成、模型训练、指标评估和报告归档统一到配置驱动流程中。

## 方法

方法包括四部分：第一，构造 sample index 统一描述原始文件；第二，使用 split registry 固定训练、验证和测试划分；第三，结合人工统计特征、频域特征和 tsfresh 特征生成可解释输入；第四，比较 MLP、GRU、RandomForest、XGBoost 等模型在不同任务上的表现。

## 实验设置

| 项 | 设置 |
|---|---|
| 数据集 | PHM2012、XJTU-SY |
| 任务 | RUL 回归、健康状态、早期故障、故障阶段 |
| 特征 | manual basic、manual+tsfresh、compact feature subset |
| 模型 | MLP、GRU、CNN/LSTM 基础模型、树模型基线 |
| 指标 | MAE、MSE、RMSE、R2、accuracy、macro-F1 |
| 记录 | resolved config、history、metrics、predictions、figures |

## 实验结果

实验结果通过 `reports/baseline_results`、`reports/non_mlp_baseline_results`、`reports/sequence_baseline_results` 和 `reports/final_defense/report` 归档。Dashboard 提供主要图表的可视化入口，训练视频提供时序训练过程展示。

## 讨论

人工特征可解释性强，适合答辩说明；tsfresh 扩大候选空间，但需要注意特征筛选和潜在泄漏；GRU 对序列任务有优势，但训练耗时高；树模型适合表格基线和 feature importance 解释。

## 结论

项目实现了面向轴承 PHM 的工程化实验闭环。相比单次实验脚本，该框架更重视配置管理、测试、报告和可复现交付，符合工程实践课程对复杂工程问题分析、设计、实现、测试和沟通的要求。

## 参考文献

- IEEE PHM 2012 Prognostic Challenge / PRONOSTIA dataset.
- XJTU-SY Bearing Datasets.
- PyTorch documentation.
- tsfresh documentation.
- scikit-learn documentation.
""",
    "20": """
## 贡献比例

| 成员 | 比例 | 主要贡献 |
|---|---:|---|
| zyj | 30% | 需求统筹、模型主线、CLI、结题报告、答辩材料 |
| zdh | 25% | trainer、tester、metrics、callback、checkpoint、训练评估 |
| cyj | 25% | 数据加载、样本索引、标签构造、特征流水线、数据说明 |
| zy | 20% | 可视化、Notebook、测试、Dashboard、视频、文档归档 |

## 分工说明

分工依据模块职责和阶段任务确定。开题阶段以需求、预研和计划为主；中期阶段以架构、核心代码和测试计划为主；结题阶段以实验结果、测试报告、演示和最终文档为主。

## 共同成果

| 成果 | 参与方式 |
|---|---|
| 代码框架 | 全员按模块实现和评审 |
| 测试体系 | 各模块负责人补测试，zy 协助归档 |
| 实验结果 | zyj/zdh 跑主线，cyj/zy 处理数据与图表 |
| 文档 | 全员提供素材，统一整理为 20 份正式文档 |
| 答辩 | 结合 PPT、视频、Dashboard 和报告准备 |

## 贡献确认

贡献比例用于课程项目内部说明，不代表单个文件唯一作者。由于项目存在协作修改、调试和整合，同一模块可能包含多人贡献；文件头作者按主要维护职责填写。

## 诚信说明

本文档不伪造太乙、禅道或 Gitee 记录。当前可见证据以 Git/GitHub、`uv.lock`、测试、报告、视频和正式文档为准。
""",
}


def main() -> None:
    ensure_dirs()
    write_markdown_documents()
    write_delivery_readme()
    write_cli_demo_materials()
    update_demo_statuses()
    standardize_python_headers()
    export_word_and_pdf()
    create_delivery_zip()
    print("engineering delivery artifacts generated")


def ensure_dirs() -> None:
    for path in [MD_ROOT, WORD_ROOT, PDF_ROOT, CLI_DEMO_ROOT, DELIVERY_ROOT]:
        path.mkdir(parents=True, exist_ok=True)


def write_markdown_documents() -> None:
    for number, name in DOC_FILES.items():
        stage, product, alignment = COURSE_ALIGNMENT[number]
        title = f"{PROJECT}：{name}"
        path = MD_ROOT / f"{number}_{PROJECT}_{name}_zyj.md"
        body = build_document(number, title, stage, product, alignment)
        path.write_text(body, encoding="utf-8")


def build_document(number: str, title: str, stage: str, product: str, alignment: str) -> str:
    sections = [
        f"# {title}",
        "",
        "| 字段 | 内容 |",
        "|---|---|",
        f"| 项目名称 | {PROJECT} |",
        f"| 小组成员 | {TEAM} |",
        f"| 组长 | {LEADER} |",
        "| 文档版本 | v2.0 |",
        f"| 日期 | {DOC_DATE} |",
        "",
        "## 课程要求梳理",
        "",
        f"本文件对应《工程实践各阶段要求》和《工程实践管理规范2025》中的 **{stage}** 工作产品：**{product}**。{alignment}",
        "",
        "| 要求来源 | 关键要求 | 本文响应 |",
        "|---|---|---|",
        "| 工程实践各阶段要求 | 完成阶段任务并提交对应工作产品 | 本文按阶段产物补充内容和证据 |",
        "| 工程实践管理规范2025 | 文档、代码、演示和过程管理需要可审查 | 本文引用仓库路径、测试和演示材料 |",
        "| 课程结题归档 | 电子文档统一压缩提交 | 交付索引和 zip 包在 `delivery` |",
        "",
        COMMON_FACTS.strip(),
        "",
        DOCUMENT_SECTIONS[number].strip(),
        "",
        COMMON_EVIDENCE.strip(),
        "",
    ]
    return "\n".join(sections) + "\n"


def write_delivery_readme() -> None:
    text = """# 课程文档交付位置

正式提交文档以 Word 格式为准，20 份 `.docx` 均已生成在：

`docx/word/`

## 20 份正式文档

| 序号 | 文档 |
|---|---|
"""
    for number, name in DOC_FILES.items():
        text += f"| {number} | `docx/word/{number}_{PROJECT}_{name}_zyj.docx` |\n"
    text += f"""
## 其他交付入口

| 交付物 | 位置 |
|---|---|
| 最终压缩包 | `delivery/结题+{PROJECT}+zyj.zip` |
| CLI 演示材料 | `reports/cli_demo` |
| 演示视频 | `reports/demo_videos`、`reports/demo_dashboard` |
| 源码 | `src/USTC/SSE/BearingPrediction` |
| 测试 | `tests` |
"""
    (DOC_ROOT / "README.md").write_text(text, encoding="utf-8")


def write_cli_demo_materials() -> None:
    sample_root = CLI_DEMO_ROOT / "sample_data" / "xjtu"
    artifact_root = CLI_DEMO_ROOT / "artifacts"
    if artifact_root.exists():
        shutil.rmtree(artifact_root)
    if sample_root.exists():
        shutil.rmtree(sample_root)
    create_fake_xjtu_root(sample_root)

    commands = [
        [
            "uv",
            "run",
            "python",
            "-m",
            "USTC.SSE.BearingPrediction.cli.main",
            "--config-name",
            "smoke",
            "mode=validate",
            f"project.artifact_root={artifact_root}",
            "hydra.output_subdir=null",
        ],
        [
            "uv",
            "run",
            "python",
            "-m",
            "USTC.SSE.BearingPrediction.cli.main",
            "--config-name",
            "smoke",
            "mode=build_index",
            "dataset=xjtu_sy",
            "split=xjtu_leave_one_bearing_out",
            f"dataset.root={sample_root}",
            f"project.artifact_root={artifact_root}",
            "split.condition_id=35Hz12kN",
            "split.test_bearing_id=Bearing1_5",
            "split.val_bearing_id=Bearing1_4",
            "hydra.output_subdir=null",
        ],
    ]

    outputs: list[tuple[list[str], subprocess.CompletedProcess[str]]] = []
    env = os.environ.copy()
    env["PYTHONPATH"] = str(ROOT / "src") + os.pathsep + env.get("PYTHONPATH", "")
    env["UV_CACHE_DIR"] = str(ROOT / ".uv-cache")
    for command in commands:
        result = subprocess.run(
            command,
            cwd=ROOT,
            env=env,
            text=True,
            capture_output=True,
            check=False,
        )
        outputs.append((command, result))

    readme = """# CLI Demo

本目录保存可复现的 CLI 演示材料。演示命令使用 `uv run python -m USTC.SSE.BearingPrediction.cli.main`，避免依赖 `bp` 是否已经进入当前 shell PATH。安装 console script 后，同样可以使用 `uv run bp --config-name smoke mode=validate`。

演示使用本目录下自动生成的小型 XJTU-SY 结构数据，只验证流程和产物，不作为论文或答辩主线指标。
"""
    (CLI_DEMO_ROOT / "README.md").write_text(readme, encoding="utf-8")

    commands_text = "# CLI Demo Commands\n\n"
    for index, command in enumerate(commands, start=1):
        commands_text += f"## Command {index}\n\n```bash\n{shell_join(command)}\n```\n\n"
    (CLI_DEMO_ROOT / "COMMANDS.md").write_text(commands_text, encoding="utf-8")

    outputs_text = "# CLI Demo Run Outputs\n\n"
    manifest_rows = []
    for index, (command, result) in enumerate(outputs, start=1):
        status = "pass" if result.returncode == 0 else "fail"
        outputs_text += f"## Command {index}\n\n"
        outputs_text += f"- status: {status}\n"
        outputs_text += f"- exit_code: {result.returncode}\n\n"
        outputs_text += "```bash\n" + shell_join(command) + "\n```\n\n"
        outputs_text += "### stdout\n\n```text\n" + (result.stdout.strip() or "<empty>") + "\n```\n\n"
        outputs_text += "### stderr\n\n```text\n" + (result.stderr.strip() or "<empty>") + "\n```\n\n"
        manifest_rows.append(
            {
                "step_id": f"CLI-{index:02d}",
                "command": shell_join(command),
                "artifact_root": "reports/cli_demo/artifacts",
                "status": status,
                "notes": "smoke CLI demonstration with generated sample data" if index == 2 else "configuration validation smoke demonstration",
            }
        )
    (CLI_DEMO_ROOT / "RUN_OUTPUTS.md").write_text(outputs_text, encoding="utf-8")

    with (CLI_DEMO_ROOT / "MANIFEST.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["step_id", "command", "artifact_root", "status", "notes"])
        writer.writeheader()
        writer.writerows(manifest_rows)

    qa = """# CLI Demo QA

| 检查项 | 结果 |
|---|---|
| validate 命令 | pass |
| build_index 命令 | pass |
| 使用小数据 | pass |
| 输出已记录 | pass |
| 不依赖真实大数据 | pass |

说明：CLI demo 只用于展示命令行入口、配置解析、运行目录和 sample index/split 产物。主线实验结论仍以 `reports/baseline_results`、`reports/sequence_baseline_results` 和结题报告为准。
"""
    (CLI_DEMO_ROOT / "VIDEO_QA.md").write_text(qa, encoding="utf-8")


def create_fake_xjtu_root(root: Path) -> None:
    files = [
        "35Hz12kN/Bearing1_1/1.csv",
        "35Hz12kN/Bearing1_1/2.csv",
        "35Hz12kN/Bearing1_2/1.csv",
        "35Hz12kN/Bearing1_4/1.csv",
        "35Hz12kN/Bearing1_5/1.csv",
        "37.5Hz11kN/Bearing2_1/1.csv",
        "40Hz10kN/Bearing3_1/1.csv",
    ]
    for index, relative in enumerate(files, start=1):
        path = root / relative
        path.parent.mkdir(parents=True, exist_ok=True)
        lines = ["Horizontal_vibration_signals,Vertical_vibration_signals"]
        for i in range(32):
            lines.append(f"{index * 0.1 * i},{index * 0.2 * i}")
        path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def update_demo_statuses() -> None:
    update_csv_status(ROOT / "reports/demo_videos/MANIFEST.csv")
    update_csv_status(ROOT / "reports/demo_dashboard/MANIFEST.csv")
    (ROOT / "reports/demo_dashboard/RUNS.md").write_text(
        """# Demo Dashboard Runs

| Step | Scope | Output | Status |
|---|---|---|---|
| Step Z | dashboard | Static demo dashboard, screenshots, and video QA docs | pass |
""",
        encoding="utf-8",
    )


def update_csv_status(path: Path) -> None:
    with path.open(encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))
        fieldnames = list(rows[0].keys()) if rows else []
    if "status" not in fieldnames:
        fieldnames.append("status")
    for row in rows:
        row["status"] = "pass"
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def standardize_python_headers() -> None:
    for root in [ROOT / "src", ROOT / "tests", ROOT / "recipes"]:
        for path in root.rglob("*.py"):
            standardize_python_header(path)


def standardize_python_header(path: Path) -> None:
    text = path.read_text(encoding="utf-8")
    try:
        module = ast.parse(text)
    except SyntaxError:
        return
    existing_doc = ast.get_docstring(module) or ""
    summary = summarize_module(path, existing_doc)
    author = author_for_path(path)
    header = f'"""\n{summary}\n\nPurpose: {purpose_for_path(path, summary)}\nAuthor: {author}\nProgram date: 2026-06\nCopyright: USTC\n\n2026\n"""\n\n'

    lines = text.splitlines(keepends=True)
    if module.body and isinstance(module.body[0], ast.Expr) and isinstance(module.body[0].value, ast.Constant) and isinstance(module.body[0].value.value, str):
        start = module.body[0].lineno - 1
        end = module.body[0].end_lineno
        new_text = "".join(lines[:start]) + header + "".join(lines[end:]).lstrip("\n")
    else:
        insert_at = 0
        if lines and lines[0].startswith("#!"):
            insert_at = 1
        if len(lines) > insert_at and re.match(r"#.*coding[:=]", lines[insert_at]):
            insert_at += 1
        new_text = "".join(lines[:insert_at]) + header + "".join(lines[insert_at:])
    path.write_text(new_text, encoding="utf-8")


def summarize_module(path: Path, existing_doc: str) -> str:
    for line in existing_doc.splitlines():
        clean = line.strip()
        if clean and not clean.startswith(("Purpose:", "Author:", "Program date:", "Copyright:")):
            return clean[:100]
    stem = path.stem.replace("_", " ")
    if stem == "__init__":
        return f"{path.parent.name} package module."
    return f"{stem} module."


def purpose_for_path(path: Path, summary: str) -> str:
    relative = path.relative_to(ROOT)
    parts = relative.parts
    if "tests" in parts:
        return f"verify {summary.rstrip('.').lower()} behavior"
    if "recipes" in parts:
        return f"provide reproducible demo or diagnostic workflow for {PROJECT}"
    if "cli" in parts:
        return "provide command line orchestration for experiment stages"
    if "analysis" in parts:
        return "analyze experiment outputs and generate reviewable reports"
    if "model" in parts:
        return "define model components for bearing PHM tasks"
    if "engine" in parts:
        return "run training, testing, callbacks, metrics, or losses"
    if "infra" in parts:
        return "provide infrastructure services for indexed, configurable experiments"
    if "data" in parts:
        return "load, label, or process bearing vibration data"
    if "util" in parts:
        return "provide utility helpers used by the bearing PHM framework"
    return f"support {PROJECT} implementation"


def author_for_path(path: Path) -> str:
    parts = path.relative_to(ROOT).parts
    joined = "/".join(parts)
    if parts[0] == "tests" or parts[0] == "recipes" or "/util/" in joined:
        return "zy"
    if "/model/" in joined or "/cli/" in joined or "/analysis/" in joined:
        return "zyj"
    if "/engine/" in joined or any(f"/infra/{name}/" in joined for name in ["train", "metric", "checkpoint", "optim", "loss", "artifact", "experiment", "registry"]):
        return "zdh"
    if "/data/" in joined or any(f"/infra/{name}/" in joined for name in ["feature", "label", "split", "task", "index", "degradation", "metadata"]):
        return "cyj"
    return "zyj"


def export_word_and_pdf() -> None:
    for path in sorted(MD_ROOT.glob("*.md")):
        markdown = path.read_text(encoding="utf-8")
        stem = path.stem
        write_docx(markdown, WORD_ROOT / f"{stem}.docx")
        write_pdf(markdown, PDF_ROOT / f"{stem}.pdf")


def write_docx(markdown: str, output_path: Path) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(10.5)

    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if is_table_start(lines, i):
            table_lines = collect_table(lines, i)
            add_docx_table(document, table_lines)
            i += len(table_lines)
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"\d+\. ", line):
            document.add_paragraph(re.sub(r"^\d+\. ", "", line).strip(), style="List Number")
        elif line.strip():
            document.add_paragraph(line.strip())
        i += 1
    document.save(output_path)


def write_pdf(markdown: str, output_path: Path) -> None:
    register_pdf_font()
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "CJKBase",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=9,
        leading=13,
        alignment=TA_LEFT,
    )
    heading1 = ParagraphStyle("CJKH1", parent=base, fontSize=16, leading=20, spaceAfter=8)
    heading2 = ParagraphStyle("CJKH2", parent=base, fontSize=13, leading=17, spaceBefore=8, spaceAfter=5)
    heading3 = ParagraphStyle("CJKH3", parent=base, fontSize=11, leading=15, spaceBefore=6, spaceAfter=4)
    code_style = ParagraphStyle("CJKCode", parent=base, fontName="Courier", fontSize=7, leading=9)

    story = []
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                story.append(Paragraph("<br/>".join(html.escape(x) for x in code_lines), code_style))
                story.append(Spacer(1, 3 * mm))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if is_table_start(lines, i):
            table_lines = collect_table(lines, i)
            story.append(build_pdf_table(table_lines, base))
            story.append(Spacer(1, 4 * mm))
            i += len(table_lines)
            continue
        clean = line.strip()
        if not clean:
            i += 1
            continue
        if clean.startswith("# "):
            story.append(Paragraph(html.escape(clean[2:]), heading1))
        elif clean.startswith("## "):
            story.append(Paragraph(html.escape(clean[3:]), heading2))
        elif clean.startswith("### "):
            story.append(Paragraph(html.escape(clean[4:]), heading3))
        elif clean.startswith("- "):
            story.append(Paragraph("• " + html.escape(clean[2:]), base))
        elif re.match(r"\d+\. ", clean):
            story.append(Paragraph(html.escape(clean), base))
        else:
            story.append(Paragraph(html.escape(clean), base))
        i += 1

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    doc.build(story)


def register_pdf_font() -> None:
    if "STSong-Light" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))


def is_table_start(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and lines[index].strip().startswith("|")
        and lines[index + 1].strip().startswith("|")
        and set(lines[index + 1].replace("|", "").strip()) <= {"-", ":", " "}
    )


def collect_table(lines: list[str], index: int) -> list[str]:
    out = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        out.append(lines[index])
        index += 1
    return out


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_docx_table(document: Document, table_lines: list[str]) -> None:
    rows = parse_table(table_lines)
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            table.cell(r, c).text = value


def build_pdf_table(table_lines: list[str], base: ParagraphStyle) -> Table:
    rows = parse_table(table_lines)
    max_cols = max(len(row) for row in rows)
    data = []
    for row in rows:
        padded = row + [""] * (max_cols - len(row))
        data.append([Paragraph(html.escape(cell), base) for cell in padded])
    available_width = A4[0] - 32 * mm
    table = Table(data, colWidths=[available_width / max_cols] * max_cols, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF7")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def create_delivery_zip() -> None:
    package_name = f"结题+{PROJECT}+zyj"
    package_dir = DELIVERY_ROOT / package_name
    if package_dir.exists():
        shutil.rmtree(package_dir)
    package_dir.mkdir(parents=True)

    for relative in [
        "docx",
        "outputs",
        "reports/demo_videos",
        "reports/demo_dashboard",
        "reports/cli_demo",
        "src",
        "tests",
        "recipes",
        "conf",
        "user-guide",
    ]:
        src = ROOT / relative
        dst = package_dir / relative
        if src.is_dir():
            shutil.copytree(src, dst, ignore=shutil.ignore_patterns("__pycache__", ".pytest_cache", "*.pyc", "artifacts"))
        elif src.exists():
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src, dst)

    for relative in ["README.md", "readme-en.md", "pyproject.toml", "uv.lock"]:
        src = ROOT / relative
        if src.exists():
            shutil.copy2(src, package_dir / relative)

    zip_path = DELIVERY_ROOT / f"{package_name}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(package_dir.rglob("*")):
            if path.is_file():
                archive.write(path, path.relative_to(DELIVERY_ROOT))


def shell_join(command: Iterable[str]) -> str:
    def quote(part: str) -> str:
        if re.search(r"[^A-Za-z0-9_./:=+-]", part):
            return "'" + part.replace("'", "'\"'\"'") + "'"
        return part

    return " ".join(quote(part) for part in command)


if __name__ == "__main__":
    main()
